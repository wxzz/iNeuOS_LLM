#把word文件转换为训练数据集

import os
import glob
import json
import pandas as pd
from docx import Document
import re
from typing import List, Dict, Any
from datasets import Dataset

class AdvancedWordToDatasetConverter:
    def __init__(self):
        self.patterns = {
            'qa_chinese': re.compile(r'^(问|问题|Q)[：:]\s*(.*?)\s*(答|答案|A)[：:]\s*(.*)$', re.MULTILINE),
            'qa_english': re.compile(r'^(Question|Q)[：:\s]\s*(.*?)\s*(Answer|A)[：:\s]\s*(.*)$', re.MULTILINE | re.IGNORECASE),
            'instruction_response': re.compile(r'^(指令|要求|Instruction)[：:]\s*(.*?)\s*(响应|回答|Response)[：:]\s*(.*)$', re.MULTILINE),
        }
    
    def process_directory(self, input_dir: str, output_file: str = "training_dataset.json") -> Dataset:
        """处理目录中的所有Word文档并生成训练数据集"""
        
        # 查找所有Word文档
        word_files = glob.glob(os.path.join(input_dir, "*.docx"))
        
        if not word_files:
            raise ValueError(f"在目录 {input_dir} 中未找到任何.docx文件")
        
        print(f"找到 {len(word_files)} 个Word文档")
        
        all_training_data = []
        
        for word_file in word_files:
            print(f"正在处理: {os.path.basename(word_file)}")
            training_data = self.convert_single_word_file(word_file)
            all_training_data.extend(training_data)
            print(f"  提取了 {len(training_data)} 条训练数据")
        
        # 保存为JSON文件
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(all_training_data, f, ensure_ascii=False, indent=2)
        
        print(f"\n总共生成 {len(all_training_data)} 条训练数据，保存到 {output_file}")
        
        # 转换为Hugging Face数据集格式
        dataset = Dataset.from_list(all_training_data)
        return dataset
    
    def convert_single_word_file(self, docx_path: str) -> List[Dict[str, Any]]:
        """转换单个Word文件为训练数据"""
        
        try:
            doc = Document(docx_path)
            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            
            # 尝试多种提取方法
            training_data = []
            
            # 方法1: 提取表格中的结构化数据
            table_data = self.extract_from_tables(doc)
            training_data.extend(table_data)
            
            # 方法2: 使用正则表达式模式提取Q&A
            pattern_data = self.extract_by_patterns(paragraphs)
            training_data.extend(pattern_data)
            
            # 方法3: 基于段落结构提取
            if len(training_data) < len(paragraphs) * 0.5:  # 如果提取不够充分
                paragraph_data = self.extract_from_paragraphs(paragraphs)
                training_data.extend(paragraph_data)
            
            # 方法4: 提取标题和内容结构
            section_data = self.extract_by_sections(doc)
            training_data.extend(section_data)
            
            # 去重
            training_data = self.deduplicate_data(training_data)
            
            return training_data
            
        except Exception as e:
            print(f"处理文件 {docx_path} 时出错: {e}")
            return []
    
    def extract_from_tables(self, doc) -> List[Dict[str, Any]]:
        """从表格中提取结构化数据"""
        training_data = []
        
        for table_idx, table in enumerate(doc.tables):
            try:
                # 尝试识别Q&A表格
                if len(table.rows) >= 2 and len(table.columns) >= 2:
                    # 检查第一行是否包含问题/答案标题
                    header_cells = [table.cell(0, i).text.strip().lower() for i in range(min(2, len(table.columns)))]
                    
                    is_qa_table = any(keyword in ' '.join(header_cells) for keyword in 
                                    ['问题', '问', 'question', '指令', 'instruction', '要求'])
                    
                    if is_qa_table:
                        # 从表格中提取Q&A对
                        for i in range(1, len(table.rows)):
                            try:
                                if len(table.columns) >= 2:
                                    instruction = table.cell(i, 0).text.strip()
                                    output = table.cell(i, 1).text.strip()
                                    if instruction and output and len(instruction) > 3 and len(output) > 5:
                                        training_data.append({
                                            "instruction": instruction,
                                            "output": output,
                                            "source": "table"
                                        })
                            except IndexError:
                                continue
                    else:
                        # 普通表格，将每行作为训练数据
                        for i in range(len(table.rows)):
                            row_texts = []
                            for j in range(len(table.columns)):
                                try:
                                    cell_text = table.cell(i, j).text.strip()
                                    if cell_text:
                                        row_texts.append(cell_text)
                                except IndexError:
                                    continue
                            
                            if row_texts:
                                row_content = " | ".join(row_texts)
                                if len(row_content) > 10:
                                    training_data.append({
                                        "instruction": f"请解释以下表格行内容: {row_content[:100]}...",
                                        "output": row_content,
                                        "source": "table_row"
                                    })
            
            except Exception as e:
                print(f"处理表格时出错: {e}")
                continue
        
        return training_data
    
    def extract_by_patterns(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """使用正则表达式模式提取Q&A对"""
        training_data = []
        
        for para in paragraphs:
            # 中文Q&A模式
            match = self.patterns['qa_chinese'].search(para)
            if match:
                training_data.append({
                    "instruction": match.group(2),
                    "output": match.group(4),
                    "source": "pattern_qa_chinese"
                })
                continue
            
            # 英文Q&A模式
            match = self.patterns['qa_english'].search(para)
            if match:
                training_data.append({
                    "instruction": match.group(2),
                    "output": match.group(4),
                    "source": "pattern_qa_english"
                })
                continue
            
            # 指令-响应模式
            match = self.patterns['instruction_response'].search(para)
            if match:
                training_data.append({
                    "instruction": match.group(2),
                    "output": match.group(4),
                    "source": "pattern_instruction"
                })
        
        return training_data
    
    def extract_from_paragraphs(self, paragraphs: List[str]) -> List[Dict[str, Any]]:
        """从段落中提取训练数据"""
        training_data = []
        
        for i, para in enumerate(paragraphs):
            if len(para) < 15:  # 跳过太短的段落
                continue
            
            # 根据段落内容生成多种指令
            instructions = self.generate_instructions(para)
            
            for instruction in instructions:
                training_data.append({
                    "instruction": instruction,
                    "output": para,
                    "source": "paragraph"
                })
        
        return training_data
    
    def generate_instructions(self, text: str) -> List[str]:
        """根据文本内容生成多种指令"""
        instructions = []
        
        # 前50个字符作为提示
        preview = text[:50] + "..." if len(text) > 50 else text
        
        # 基础指令模板
        base_instructions = [
            f"请详细说明：{preview}",
            f"请解释以下内容：{preview}",
            f"关于'{preview}'，请提供详细信息",
            f"总结以下内容：{preview}",
            f"请分析：{preview}",
            f"请描述：{preview}",
            f"什么是{preview}？",
            f"{preview}是什么意思？"
        ]
        
        # 根据文本特征添加特定指令
        if '?' in text or '？' in text:
            base_instructions.extend([
                f"请回答：{preview}",
                f"如何解决：{preview}"
            ])
        
        if len(text) > 200:
            base_instructions.extend([
                f"请详细阐述：{preview}",
                f"请全面介绍：{preview}"
            ])
        else:
            base_instructions.extend([
                f"请简要说明：{preview}",
                f"请概括：{preview}"
            ])
        
        # 选择3-5个不同的指令
        import random
        selected_instructions = random.sample(base_instructions, min(4, len(base_instructions)))
        
        return selected_instructions
    
    def extract_by_sections(self, doc) -> List[Dict[str, Any]]:
        """基于标题和内容结构提取"""
        training_data = []
        sections = []
        current_section = {"title": "", "content": []}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测标题（基于样式或文本特征）
            is_heading = (
                para.style.name.startswith('Heading') or 
                para.style.name.startswith('标题') or
                len(text) < 100 and (text.endswith('：') or not text.endswith('。')) or
                self.is_likely_heading(text)
            )
            
            if is_heading and current_section["content"]:
                # 保存当前章节
                sections.append(current_section)
                current_section = {"title": text, "content": []}
            elif is_heading:
                # 新章节开始
                current_section["title"] = text
            else:
                # 添加到当前章节内容
                current_section["content"].append(text)
        
        # 添加最后一个章节
        if current_section["content"]:
            sections.append(current_section)
        
        # 为每个章节创建训练数据
        for section in sections:
            if section["title"] and section["content"]:
                full_content = " ".join(section["content"])
                if len(full_content) > 20:  # 确保内容足够长
                    instructions = self.generate_section_instructions(section["title"])
                    for instruction in instructions:
                        training_data.append({
                            "instruction": instruction,
                            "output": full_content,
                            "source": "section"
                        })
        
        return training_data
    
    def is_likely_heading(self, text: str) -> bool:
        """判断文本是否可能是标题"""
        if len(text) > 150:
            return False
        
        heading_indicators = [
            '第一章', '第二章', '第一节', '第二节', '一、', '二、', '三、',
            '1. ', '2. ', '3. ', '(1)', '(2)', '(3)',
            '第1章', '第2章', '第3章', '第一部分', '第二部分'
        ]
        
        return any(indicator in text for indicator in heading_indicators)
    
    def generate_section_instructions(self, title: str) -> List[str]:
        """为章节标题生成指令"""
        instructions = [
            f"请介绍{title}",
            f"什么是{title}？",
            f"请详细说明{title}的相关内容",
            f"关于{title}，请提供详细信息",
            f"总结{title}的主要内容",
            f"{title}的概念是什么？",
            f"请解释{title}",
            f"{title}包括哪些内容？"
        ]
        
        import random
        return random.sample(instructions, min(4, len(instructions)))
    
    def deduplicate_data(self, training_data: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """去重训练数据"""
        seen = set()
        deduplicated = []
        
        for item in training_data:
            # 基于instruction和output的组合创建唯一标识
            key = (item["instruction"][:100], item["output"][:100])
            
            if key not in seen:
                seen.add(key)
                deduplicated.append(item)
        
        print(f"去重后数据: {len(training_data)} -> {len(deduplicated)}")
        return deduplicated

def analyze_dataset(dataset_path: str):
    """分析生成的数据集"""
    with open(dataset_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    print(f"\n数据集分析:")
    print(f"总样本数: {len(data)}")
    
    # 统计不同来源的数据
    sources = {}
    for item in data:
        source = item.get("source", "unknown")
        sources[source] = sources.get(source, 0) + 1
    
    print("数据来源分布:")
    for source, count in sources.items():
        print(f"  {source}: {count}")
    
    # 统计指令和输出的平均长度
    inst_lengths = [len(item["instruction"]) for item in data]
    output_lengths = [len(item["output"]) for item in data]
    
    print(f"指令平均长度: {sum(inst_lengths) / len(inst_lengths):.1f} 字符")
    print(f"输出平均长度: {sum(output_lengths) / len(output_lengths):.1f} 字符")
    
    # 显示前几个样本
    print("\n前3个样本示例:")
    for i, item in enumerate(data[:3]):
        print(f"样本 {i+1}:")
        print(f"  指令: {item['instruction']}")
        print(f"  输出: {item['output'][:100]}...")
        print(f"  来源: {item.get('source', 'unknown')}")
        print()

# 主执行函数
def main():
    input_dir = "docx_input"
    output_file = "training_dataset.json"
    
    # 检查输入目录是否存在
    if not os.path.exists(input_dir):
        print(f"错误: 目录 '{input_dir}' 不存在")
        return
    
    # 创建转换器并处理目录
    converter = AdvancedWordToDatasetConverter()
    
    try:
        dataset = converter.process_directory(input_dir, output_file)
        
        # 分析生成的数据集
        analyze_dataset(output_file)
        
        print("转换完成！数据集已保存为 training_dataset.json")
        print("现在你可以使用这个文件来训练你的模型了。")
        
    except Exception as e:
        print(f"处理过程中出错: {e}")

if __name__ == "__main__":
    main()